"""Unified document generator. Builds PDF and DOCX from structured content blocks.

Block types are dataclasses. Each output format has a Renderer that knows how
to render every block type. Adding a new block = add a dataclass + update
renderers. Adding a new format = add a Renderer class.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol


# ---------------------------------------------------------------------------
# Block dataclasses
# ---------------------------------------------------------------------------


@dataclass
class Heading:
    text: str
    level: int = 1


@dataclass
class Text:
    text: str


@dataclass
class ListBlock:
    items: list[str]
    ordered: bool = False


@dataclass
class Table:
    headers: list[str]
    rows: list[list]


@dataclass
class Divider:
    pass


@dataclass
class PageBreak:
    pass


Block = Heading | Text | ListBlock | Table | Divider | PageBreak


# ---------------------------------------------------------------------------
# Style configuration (injectable, no renderer changes needed to restyle)
# ---------------------------------------------------------------------------


@dataclass
class StyleConfig:
    font_body: str = "Calibri"
    font_heading: str = "Calibri"
    font_mono: str = "Courier New"
    font_size_body: int = 11
    font_size_h1: int = 22
    font_size_h2: int = 16
    font_size_h3: int = 13
    color_subtitle: str = "#666666"
    color_author: str = "#999999"
    color_divider: str = "#cccccc"
    color_table_header_bg: str = "#f0f0f0"
    color_table_grid: str = "#cccccc"


def parse_block(raw: dict) -> Block:
    """Parse a raw dict into a typed Block."""
    btype = raw.get("type", "text")
    match btype:
        case "heading":
            return Heading(raw.get("text", ""), raw.get("level", 1))
        case "text":
            return Text(raw.get("text", ""))
        case "list":
            return ListBlock(raw.get("items", []), raw.get("ordered", False))
        case "table":
            return Table(raw.get("headers", []), raw.get("rows", []))
        case "divider":
            return Divider()
        case "page_break":
            return PageBreak()
        case _:
            raise ValueError(f"Unknown block type: {btype}")


# ---------------------------------------------------------------------------
# Renderer protocol
# ---------------------------------------------------------------------------


class Renderer(Protocol):
    def begin(self, title: str | None, subtitle: str | None, author: str | None) -> None: ...
    def render(self, block: Block) -> None: ...
    def save(self, path: str) -> None: ...


# ---------------------------------------------------------------------------
# PDF renderer
# ---------------------------------------------------------------------------


class PdfRenderer:
    def __init__(self, style: StyleConfig | None = None):
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch

        self._cfg = style or StyleConfig()
        self._pagesize = letter
        self._margins = {
            "leftMargin": inch, "rightMargin": inch,
            "topMargin": inch, "bottomMargin": inch,
        }
        self._styles = _build_pdf_styles(self._cfg)
        self._story: list = []

    def begin(self, title, subtitle, author):
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.units import inch

        if title:
            self._story.append(Paragraph(title, self._styles["doc_title"]))
        if subtitle:
            self._story.append(Spacer(1, 0.1 * inch))
            self._story.append(Paragraph(subtitle, self._styles["subtitle"]))
        if author:
            self._story.append(Spacer(1, 0.05 * inch))
            self._story.append(Paragraph(author, self._styles["author"]))
        if title:
            self._story.append(Spacer(1, 0.3 * inch))

    def render(self, block: Block) -> None:
        from reportlab.platypus import Paragraph, Spacer, PageBreak as PdfPageBreak, HRFlowable
        from reportlab.lib.units import inch

        match block:
            case Heading(text, level):
                key = f"h{min(level, 3)}"
                self._story.append(Paragraph(text, self._styles[key]))
            case Text(text):
                parts = []
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        parts.append(Paragraph(line, self._styles["body"]))
                if parts:
                    self._story.extend(parts)
                else:
                    self._story.append(Spacer(1, 0.05 * inch))
            case ListBlock(items, ordered):
                for i, item in enumerate(items):
                    prefix = f"{i + 1}. " if ordered else "- "
                    self._story.append(
                        Paragraph(f"{prefix}{item}", self._styles["list_item"])
                    )
            case Table(headers, rows):
                self._story.append(Spacer(1, 0.05 * inch))
                self._story.append(_build_pdf_table(headers, rows, self._cfg))
                self._story.append(Spacer(1, 0.1 * inch))
            case Divider():
                self._story.append(Spacer(1, 0.1 * inch))
                self._story.append(HRFlowable(width="100%"))
                self._story.append(Spacer(1, 0.1 * inch))
            case PageBreak():
                self._story.append(PdfPageBreak())

    def save(self, path: str) -> None:
        from reportlab.platypus import SimpleDocTemplate, Spacer
        from reportlab.lib.units import inch

        if not self._story:
            self._story.append(Spacer(1, 0.1 * inch))

        pdf = SimpleDocTemplate(path, pagesize=self._pagesize, **self._margins)
        pdf.build(self._story)


def _build_pdf_styles(cfg: StyleConfig):
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.colors import HexColor

    base = getSampleStyleSheet()
    return {
        "doc_title": ParagraphStyle(
            "DocTitle", parent=base["Title"], fontSize=cfg.font_size_h1, spaceAfter=4,
        ),
        "subtitle": ParagraphStyle(
            "Subtitle", parent=base["Normal"], fontSize=14,
            alignment=TA_CENTER, textColor=HexColor(cfg.color_subtitle),
        ),
        "author": ParagraphStyle(
            "Author", parent=base["Normal"], fontSize=cfg.font_size_body,
            alignment=TA_CENTER, textColor=HexColor(cfg.color_author),
        ),
        "h1": ParagraphStyle(
            "H1", parent=base["Heading1"], fontSize=cfg.font_size_h2,
            spaceBefore=16, spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "H2", parent=base["Heading2"], fontSize=cfg.font_size_h3,
            spaceBefore=12, spaceAfter=6,
        ),
        "h3": ParagraphStyle(
            "H3", parent=base["Heading3"], fontSize=cfg.font_size_body,
            spaceBefore=10, spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "Body", parent=base["Normal"], fontSize=cfg.font_size_body - 1,
            leading=14, spaceAfter=6,
        ),
        "list_item": ParagraphStyle(
            "ListItem", parent=base["Normal"], fontSize=cfg.font_size_body - 1,
            leading=14, leftIndent=20, spaceAfter=3,
        ),
    }


def _build_pdf_table(headers, rows, cfg: StyleConfig):
    from reportlab.platypus import Table as RlTable, TableStyle
    from reportlab.lib.colors import HexColor

    data = [headers] + rows if headers else rows
    if not data:
        data = [[""]]
    t = RlTable(data)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(cfg.color_table_header_bg)),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor(cfg.color_table_grid)),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    return t


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------


class DocxRenderer:
    def __init__(self, style: StyleConfig | None = None):
        from docx import Document
        from docx.shared import Pt

        self._cfg = style or StyleConfig()
        self._doc = Document()
        normal = self._doc.styles["Normal"]
        normal.font.size = Pt(self._cfg.font_size_body)
        normal.font.name = self._cfg.font_body

    def begin(self, title, subtitle, author):
        from docx.shared import Pt
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if title:
            self._doc.add_heading(title, level=0)
        if subtitle:
            p = self._doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = _hex_to_rgb(self._cfg.color_subtitle)
        if author:
            p = self._doc.add_paragraph(author)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = _hex_to_rgb(self._cfg.color_author)
                run.font.size = Pt(self._cfg.font_size_body - 1)

    def render(self, block: Block) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        match block:
            case Heading(text, level):
                self._doc.add_heading(text, level=min(level, 4))
            case Text(text):
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        self._doc.add_paragraph(line)
            case ListBlock(items, ordered):
                style = "List Number" if ordered else "List Bullet"
                for item in items:
                    self._doc.add_paragraph(item, style=style)
            case Table(headers, rows):
                cols = len(headers) if headers else (len(rows[0]) if rows else 1)
                total_rows = (1 if headers else 0) + len(rows)
                if total_rows == 0:
                    return
                table = self._doc.add_table(rows=total_rows, cols=cols)
                table.style = "Table Grid"
                row_idx = 0
                if headers:
                    for j, h in enumerate(headers):
                        cell = table.cell(0, j)
                        cell.text = str(h)
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                    row_idx = 1
                for data_row in rows:
                    for j, val in enumerate(data_row):
                        table.cell(row_idx, j).text = str(val)
                    row_idx += 1
            case Divider():
                p = self._doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("_" * 50)
                run.font.color.rgb = _hex_to_rgb(self._cfg.color_divider)
            case PageBreak():
                self._doc.add_page_break()

    def save(self, path: str) -> None:
        self._doc.save(path)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _hex_to_rgb(hex_color: str):
    """Convert '#RRGGBB' to docx RGBColor."""
    from docx.shared import RGBColor
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


_RENDERERS: dict[str, type] = {
    ".pdf": PdfRenderer,
    ".docx": DocxRenderer,
}


def _render_doc(renderer: Renderer, doc: dict, path: str) -> None:
    renderer.begin(doc.get("title"), doc.get("subtitle"), doc.get("author"))
    for raw in doc.get("content", []):
        renderer.render(parse_block(raw))
    renderer.save(path)


def generate_pdf(path: str, doc: dict, style: StyleConfig | None = None) -> None:
    _render_doc(PdfRenderer(style), doc, path)


def generate_docx(path: str, doc: dict, style: StyleConfig | None = None) -> None:
    _render_doc(DocxRenderer(style), doc, path)


def generate(path: str, doc: dict, style: StyleConfig | None = None) -> None:
    """Auto-detect format from file extension and generate."""
    ext = Path(path).suffix.lower()
    cls = _RENDERERS.get(ext)
    if not cls:
        raise ValueError(f"Unsupported format: {ext}")
    _render_doc(cls(style), doc, path)
