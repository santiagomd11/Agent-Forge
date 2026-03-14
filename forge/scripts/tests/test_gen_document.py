"""Tests for unified document generator (gen_document.py)."""

import os
import tempfile

import pytest
from docx import Document as DocxDocument

from forge.scripts.src.gen_document import (
    Heading,
    Text,
    ListBlock,
    Table,
    Divider,
    PageBreak,
    StyleConfig,
    parse_block,
    generate_pdf,
    generate_docx,
    generate,
    PdfRenderer,
    DocxRenderer,
)


# --- Fixtures ---


@pytest.fixture
def pdf_path():
    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    yield path
    if os.path.exists(path):
        os.unlink(path)


@pytest.fixture
def docx_path():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    yield path
    if os.path.exists(path):
        os.unlink(path)


# --- Block dataclass tests ---


class TestBlockClasses:

    def test_heading_defaults(self):
        h = Heading("Title")
        assert h.text == "Title"
        assert h.level == 1

    def test_heading_custom_level(self):
        h = Heading("Sub", level=3)
        assert h.level == 3

    def test_text_block(self):
        t = Text("Hello world")
        assert t.text == "Hello world"

    def test_list_block_defaults(self):
        lb = ListBlock(["a", "b"])
        assert lb.items == ["a", "b"]
        assert lb.ordered is False

    def test_list_block_ordered(self):
        lb = ListBlock(["x"], ordered=True)
        assert lb.ordered is True

    def test_table_block(self):
        t = Table(["Name"], [["Alice"]])
        assert t.headers == ["Name"]
        assert t.rows == [["Alice"]]

    def test_divider(self):
        d = Divider()
        assert isinstance(d, Divider)

    def test_page_break(self):
        pb = PageBreak()
        assert isinstance(pb, PageBreak)


# --- parse_block tests ---


class TestParseBlock:

    def test_parse_heading(self):
        b = parse_block({"type": "heading", "text": "Hi", "level": 2})
        assert isinstance(b, Heading)
        assert b.text == "Hi"
        assert b.level == 2

    def test_parse_text(self):
        b = parse_block({"type": "text", "text": "Hello"})
        assert isinstance(b, Text)
        assert b.text == "Hello"

    def test_parse_text_default(self):
        b = parse_block({"text": "No type"})
        assert isinstance(b, Text)

    def test_parse_list(self):
        b = parse_block({"type": "list", "items": ["a"], "ordered": True})
        assert isinstance(b, ListBlock)
        assert b.items == ["a"]
        assert b.ordered is True

    def test_parse_table(self):
        b = parse_block({"type": "table", "headers": ["H"], "rows": [["1"]]})
        assert isinstance(b, Table)

    def test_parse_divider(self):
        b = parse_block({"type": "divider"})
        assert isinstance(b, Divider)

    def test_parse_page_break(self):
        b = parse_block({"type": "page_break"})
        assert isinstance(b, PageBreak)

    def test_parse_unknown_raises(self):
        with pytest.raises(ValueError, match="Unknown block type"):
            parse_block({"type": "sparkle"})


# --- Renderer instantiation ---


class TestRenderers:

    def test_pdf_renderer_creates(self):
        r = PdfRenderer()
        assert r is not None

    def test_docx_renderer_creates(self):
        r = DocxRenderer()
        assert r is not None


# --- PDF generation via generate_pdf ---


class TestGeneratePdf:

    def test_creates_file(self, pdf_path):
        generate_pdf(pdf_path, {"content": [{"type": "text", "text": "Hello"}]})
        assert os.path.isfile(pdf_path)

    def test_valid_pdf_header(self, pdf_path):
        generate_pdf(pdf_path, {"content": [{"type": "text", "text": "Check"}]})
        with open(pdf_path, "rb") as f:
            assert f.read(5) == b"%PDF-"

    def test_title_and_subtitle(self, pdf_path):
        generate_pdf(pdf_path, {
            "title": "Quarterly Report",
            "subtitle": "Q1 2026",
            "author": "Sales Team",
            "content": [{"type": "text", "text": "Revenue grew 15%."}],
        })
        assert os.path.getsize(pdf_path) > 0

    def test_headings(self, pdf_path):
        generate_pdf(pdf_path, {"content": [
            {"type": "heading", "text": "H1", "level": 1},
            {"type": "heading", "text": "H2", "level": 2},
            {"type": "heading", "text": "H3", "level": 3},
        ]})
        assert os.path.getsize(pdf_path) > 0

    def test_unordered_list(self, pdf_path):
        generate_pdf(pdf_path, {"content": [
            {"type": "list", "items": ["Apples", "Bananas"]},
        ]})
        assert os.path.getsize(pdf_path) > 0

    def test_ordered_list(self, pdf_path):
        generate_pdf(pdf_path, {"content": [
            {"type": "list", "items": ["First", "Second"], "ordered": True},
        ]})
        assert os.path.getsize(pdf_path) > 0

    def test_table(self, pdf_path):
        generate_pdf(pdf_path, {"content": [
            {"type": "table", "headers": ["Name", "Score"], "rows": [["Alice", "95"]]},
        ]})
        assert os.path.getsize(pdf_path) > 0

    def test_divider(self, pdf_path):
        generate_pdf(pdf_path, {"content": [
            {"type": "text", "text": "Before"},
            {"type": "divider"},
            {"type": "text", "text": "After"},
        ]})
        assert os.path.getsize(pdf_path) > 0

    def test_page_break(self, pdf_path):
        generate_pdf(pdf_path, {"content": [
            {"type": "text", "text": "Page 1"},
            {"type": "page_break"},
            {"type": "text", "text": "Page 2"},
        ]})
        assert os.path.getsize(pdf_path) > 0

    def test_empty_doc(self, pdf_path):
        generate_pdf(pdf_path, {})
        assert os.path.isfile(pdf_path)

    def test_full_document(self, pdf_path):
        generate_pdf(pdf_path, {
            "title": "Analysis Report",
            "subtitle": "March 2026",
            "author": "Research Agent",
            "content": [
                {"type": "heading", "text": "Summary", "level": 1},
                {"type": "text", "text": "Q1 data analysis."},
                {"type": "list", "items": ["Revenue up 12%", "Costs down 3%"]},
                {"type": "divider"},
                {"type": "heading", "text": "Data", "level": 2},
                {"type": "table", "headers": ["Metric", "Value"], "rows": [["Revenue", "$1.2M"]]},
                {"type": "page_break"},
                {"type": "heading", "text": "Conclusion", "level": 1},
                {"type": "text", "text": "Strong quarter."},
            ],
        })
        assert os.path.getsize(pdf_path) > 1000


# --- DOCX generation via generate_docx ---


def _docx_text(path):
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs)


class TestGenerateDocx:

    def test_creates_file(self, docx_path):
        generate_docx(docx_path, {"content": [{"type": "text", "text": "Hello"}]})
        assert os.path.isfile(docx_path)

    def test_valid_docx(self, docx_path):
        generate_docx(docx_path, {"content": [{"type": "text", "text": "Check"}]})
        doc = DocxDocument(docx_path)
        assert len(doc.paragraphs) > 0

    def test_title(self, docx_path):
        generate_docx(docx_path, {"title": "My Report", "content": []})
        assert "My Report" in _docx_text(docx_path)

    def test_subtitle(self, docx_path):
        generate_docx(docx_path, {"title": "T", "subtitle": "Sub", "content": []})
        assert "Sub" in _docx_text(docx_path)

    def test_text_content(self, docx_path):
        generate_docx(docx_path, {"content": [
            {"type": "text", "text": "Para one.\nPara two."},
        ]})
        text = _docx_text(docx_path)
        assert "Para one" in text
        assert "Para two" in text

    def test_headings(self, docx_path):
        generate_docx(docx_path, {"content": [
            {"type": "heading", "text": "Section A", "level": 1},
            {"type": "heading", "text": "Subsection", "level": 2},
        ]})
        text = _docx_text(docx_path)
        assert "Section A" in text

    def test_unordered_list(self, docx_path):
        generate_docx(docx_path, {"content": [
            {"type": "list", "items": ["Alpha", "Beta", "Gamma"]},
        ]})
        text = _docx_text(docx_path)
        assert "Alpha" in text
        assert "Gamma" in text

    def test_ordered_list(self, docx_path):
        generate_docx(docx_path, {"content": [
            {"type": "list", "items": ["First", "Second"], "ordered": True},
        ]})
        assert "First" in _docx_text(docx_path)

    def test_table(self, docx_path):
        generate_docx(docx_path, {"content": [
            {"type": "table", "headers": ["Name", "Age"], "rows": [["Alice", "30"]]},
        ]})
        doc = DocxDocument(docx_path)
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "Name"
        assert doc.tables[0].cell(1, 0).text == "Alice"

    def test_page_break(self, docx_path):
        generate_docx(docx_path, {"content": [
            {"type": "text", "text": "Page 1"},
            {"type": "page_break"},
            {"type": "text", "text": "Page 2"},
        ]})
        assert os.path.getsize(docx_path) > 0

    def test_empty_doc(self, docx_path):
        generate_docx(docx_path, {})
        assert os.path.isfile(docx_path)

    def test_full_document(self, docx_path):
        generate_docx(docx_path, {
            "title": "Test Report",
            "subtitle": "March 2026",
            "author": "Agent",
            "content": [
                {"type": "heading", "text": "Intro", "level": 1},
                {"type": "text", "text": "Overview of findings."},
                {"type": "list", "items": ["Point A", "Point B"]},
                {"type": "table", "headers": ["K", "V"], "rows": [["x", "1"]]},
                {"type": "divider"},
                {"type": "heading", "text": "End", "level": 1},
            ],
        })
        text = _docx_text(docx_path)
        assert "Test Report" in text
        assert "Overview" in text


# --- Auto-detect generate() ---


class TestAutoDetectGenerate:

    def test_auto_pdf(self, pdf_path):
        generate(pdf_path, {"content": [{"type": "text", "text": "Auto PDF"}]})
        with open(pdf_path, "rb") as f:
            assert f.read(5) == b"%PDF-"

    def test_auto_docx(self, docx_path):
        generate(docx_path, {"content": [{"type": "text", "text": "Auto DOCX"}]})
        doc = DocxDocument(docx_path)
        assert len(doc.paragraphs) > 0

    def test_unsupported_format_raises(self, tmp_path):
        bad_path = str(tmp_path / "file.xyz")
        with pytest.raises(ValueError, match="Unsupported format"):
            generate(bad_path, {"content": []})


# --- StyleConfig tests ---


class TestStyleConfig:

    def test_default_values(self):
        s = StyleConfig()
        assert s.font_body == "Calibri"
        assert s.font_heading == "Calibri"
        assert s.font_mono == "Courier New"
        assert s.font_size_body == 11
        assert s.font_size_h1 == 22
        assert s.font_size_h2 == 16
        assert s.font_size_h3 == 13
        assert s.color_subtitle == "#666666"
        assert s.color_author == "#999999"
        assert s.color_divider == "#cccccc"
        assert s.color_table_header_bg == "#f0f0f0"
        assert s.color_table_grid == "#cccccc"

    def test_custom_values(self):
        s = StyleConfig(font_body="Arial", font_size_body=14, color_subtitle="#111111")
        assert s.font_body == "Arial"
        assert s.font_size_body == 14
        assert s.color_subtitle == "#111111"
        # defaults still hold for unset fields
        assert s.font_heading == "Calibri"

    def test_pdf_with_custom_style(self, pdf_path):
        style = StyleConfig(font_body="Helvetica", font_size_body=12)
        generate_pdf(pdf_path, {
            "title": "Custom",
            "content": [{"type": "text", "text": "Styled text"}],
        }, style=style)
        assert os.path.isfile(pdf_path)
        with open(pdf_path, "rb") as f:
            assert f.read(5) == b"%PDF-"

    def test_docx_with_custom_style(self, docx_path):
        style = StyleConfig(font_body="Arial", font_size_body=14)
        generate_docx(docx_path, {
            "content": [{"type": "text", "text": "Custom font"}],
        }, style=style)
        doc = DocxDocument(docx_path)
        normal = doc.styles["Normal"]
        assert normal.font.name == "Arial"
        assert normal.font.size.pt == 14

    def test_generate_auto_with_style(self, pdf_path):
        style = StyleConfig(font_body="Times New Roman")
        generate(pdf_path, {
            "content": [{"type": "text", "text": "Auto styled"}],
        }, style=style)
        assert os.path.isfile(pdf_path)

    def test_default_style_when_none(self, docx_path):
        generate_docx(docx_path, {"content": [{"type": "text", "text": "Default"}]})
        doc = DocxDocument(docx_path)
        normal = doc.styles["Normal"]
        assert normal.font.name == "Calibri"
        assert normal.font.size.pt == 11
